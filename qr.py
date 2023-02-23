import os
from PIL import Image
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.shared import Inches
from docx2pdf import convert
import re
from generate_qr import generate_qr_with_label, generate_qr_from_file

#------------------------------------------------------------------------------------------------------------------#
#------------------------ GENERATE QR CODE FOR THE SET RANGES USING THE CUSTOM FUNCTION ---------------------------#
#------------------------------------------------------------------------------------------------------------------#
prefix = 'AEDCBD00'
start = 11420
end = 11801
step = 1
img = generate_qr_with_label(prefix, start, end, step)

#------------------------------------------------------------------------------------------------------------------#
#---------- GENERATE QR CODE FOR THE RANGE OF NUMBERS IN A CSV OR EXCEL FILE USING THE CUSTOM FUNCTION ------------#
#------------------------------------------------------------------------------------------------------------------#

file_path = os.path.join(os.path.expanduser("~"),"Downloads", "missing_qr_numbers")
# Call the generate_qr_from_file function
qr_codes = generate_qr_from_file(prefix, file_path)

orientation = "PORTRAIT" #LANDSCAPE, PORTRAIT
columns = 3

#-------------------------------------------------------------------------------------------------------------#
#-------------------------------- MERGE GENERATED CODES AND LOCATION LOGO ------------------------------------#
#-------------------------------------------------------------------------------------------------------------#
if prefix.startswith('AEDC'):
    location = "AEDC"
    disco_logo = 'aedc-logo.png'
elif prefix.startswith('ECG'):
    location = 'ECG'
    disco_logo = 'ecg-ghana.jpg'
else:
    raise ValueError('Invalid prefix')

asset = "BD"
dir_name = os.path.join(os.path.expanduser("~"),"OneDrive", "dev", "qr_code_tag_generator", "QRCodes", location)
image_files = [f for f in os.listdir(dir_name) if f.endswith(".png")] # get a list of all PNG files in the folder

global vtags_dir
vtags_dir = "VTags"

logo = os.path.join(os.getcwd(), disco_logo)
logo = Image.open(logo)
logo = logo.resize((332, 332))
logo_size = logo.size

for qc_name in image_files:
    qc = os.path.join(dir_name, qc_name)
    qc = Image.open(qc)
    qc_size = qc.size
    s2u = [i*2 for i in qc_size]
    qc = qc.resize(s2u)
    qc_size = qc.size

    logo_qc = Image.new('RGB',(332+290, 332), (250,250,250))
    logo_qc.paste(logo,(0,0))
    logo_qc.paste(qc,(332,0))

    if not os.path.exists(vtags_dir):
        os.makedirs(vtags_dir)

    logo_qc.save(os.path.join(vtags_dir, qc_name),"PNG")
    # logo_qc.show()

#-------------------------------------------------------------------------------------------------#
#------------------------ LOAD FILES INTO DOCX FILE AND CONVERT TO PDF ---------------------------#
#-------------------------------------------------------------------------------------------------#

pics = sorted(os.listdir(dir_name))
pics[:4], len(pics)

document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    if orientation == 'PORTRAIT':
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
    else:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(columns))
    cols.set(qn('w:space'), '10')

logo = os.path.join(os.getcwd(), disco_logo)

# Extract the beginning and end of the VTag range from the file names
vtags = [int(re.search(r'\d+', f).group()) for f in pics]
vtags_begin = min(vtags)
vtags_end = max(vtags)

for qrc in pics:
    paragraph = document.add_paragraph()
    pth = os.path.join(os.getcwd(), dir_name, qrc)
    qc = Image.open(pth)
    qc_size = qc.size
    s2u = [i*2 for i in qc_size]
    qc = qc.resize(s2u)
    qc_size = qc.size

    logo = os.path.join(os.getcwd(), disco_logo)
    logo = Image.open(logo)
    logo = logo.resize((332, 332))
    logo_size = logo.size

    logo_qc = Image.new('RGB',(332+290, 332), (250,250,250))
    logo_qc.paste(logo,(0,0))
    logo_qc.paste(qc,(332,0))

    # Construct the file name for the QR code image with the VTag range included
    qc_name = re.sub(r'[:\\]', '_', qrc)
    qc_range = f"{vtags_begin:05} - {vtags_end:05}"
    fpth = f"{vtags_dir}/{location} {vtags_dir} {asset} {qc_range} {qc_name}"
    logo_qc.save(fpth,"PNG")
    lqpth = os.path.join(os.getcwd(), fpth)

    run_2 = paragraph.add_run()
    run_2.add_picture(lqpth, height=Inches(1.25))

document_dir = "generated"
if not os.path.exists(document_dir):
    os.makedirs(document_dir)

docx_file_name = f'{location} {vtags_dir} {asset} {qc_range}.docx'
docx_file_name = re.sub(r'[:\\]', '_', docx_file_name)
docx_file_path = os.path.join(document_dir, docx_file_name)

# Check if the file already exists and update the filename if necessary
i = 1
while os.path.exists(docx_file_path):
    docx_file_name = f'{location} {vtags_dir} {asset} {qc_range} ({i}).docx'
    docx_file_name = re.sub(r'[:\\]', '_', docx_file_name)
    docx_file_path = os.path.join(document_dir, docx_file_name)
    i += 1

# Save the Word document with the updated filename
document.save(docx_file_path)

# Convert the Word document to PDF
pdf_file_name = f'{location} {vtags_dir} {asset} {qc_range}.pdf'
pdf_file_name = re.sub(r'[:\\]', '_', pdf_file_name)
pdf_dir_path = os.path.join(document_dir, 'pdf')
if not os.path.exists(pdf_dir_path):
    os.makedirs(pdf_dir_path)
pdf_file_path = os.path.join(pdf_dir_path, pdf_file_name)

# Check if the file already exists and update the name if necessary
if os.path.isfile(pdf_file_path):
    i = 1
    while True:
        new_file_name = f'{location} {vtags_dir} {asset} {qc_range} ({i}).pdf'
        new_file_path = os.path.join(pdf_dir_path, new_file_name)
        if not os.path.isfile(new_file_path):
            pdf_file_path = new_file_path
            break
        i += 1

convert(docx_file_path, pdf_file_path)


# Delete all image files in the QRCodes and VTags folders and their subdirectories
for root, dirs, files in os.walk("QRCodes"):
    for file in files:
        if file.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif')):
            os.remove(os.path.join(root, file))
            
for root, dirs, files in os.walk("VTags"):
    for file in files:
        if file.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif')):
            os.remove(os.path.join(root, file))
